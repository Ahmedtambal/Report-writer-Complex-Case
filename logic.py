import streamlit as st  # ✅ Correct importimport openai
from docx import Document
from datetime import datetime
import pdfplumber
import json
import os
import re  # Ensure this is included
import openai
import pytesseract
from PIL import Image
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import platform
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import mammoth
from io import StringIO


# At the top of your file with other imports
# Initialize client once at the module level
try:
    import streamlit as st
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except (AttributeError, ModuleNotFoundError):
    import os
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    if not client.api_key:
        raise ValueError("OpenAI API key not found in environment variables or Streamlit secrets")

UPLOAD_FOLDER = "uploaded_docs"  # Ensure it's defined globally


# Specify the full path to Tesseract executable if not in PATH
# Add this at the top of your file (after imports):

# Set Tesseract path based on OS
if platform.system() == 'Windows':
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
else:  # For Linux (Streamlit Cloud) and macOS
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

def clean_json_response(response_str) -> str:
    """Strip out code fences, extra markdown, etc."""
    cleaned = re.sub(r'^```json\s*|\s*```$', '', response_str, flags=re.DOTALL)
    return cleaned.strip()

def parse_json_response(response_content: str, context: str = "") -> dict:
    """Parse and validate JSON response with error handling"""
    try:
        return json.loads(response_content)
    except json.JSONDecodeError as e:
        error_msg = (
            f"JSON parsing failed in {context}:\n"
            f"Error: {str(e)}\n"
            f"Original content: {repr(response_content)}\n"
        )
        raise ValueError(error_msg) from e
    
def extract_text_from_file(file_path):
    """Extract text from PDF, docx, or image (png/jpg/jpeg)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return extract_text_from_image(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return ""
    
def extract_text_from_image(image_path):
    """
    Extract text from an image using Tesseract OCR.
    
    Args:
    - image_path (str): Path to the image file.

    Returns:
    - str: Extracted text from the image.
    """
    # Load the image
    image = Image.open(image_path)

    # Extract text using Tesseract OCR
    extracted_text = pytesseract.image_to_string(image)

    return extracted_text

# logic.py (snippet)

def extract_text_from_docx(file_path):
    with open(file_path, "rb") as docx_file:
        # Convert to Markdown (you can also do .convert_to_html)
        result = mammoth.convert_to_markdown(docx_file)
        text = result.value  # The generated Markdown
    # You could remove Markdown formatting here, or just keep it
    return text


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file using pdfplumber.
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages])
        return text
    except Exception as e:
        raise ValueError("Error extracting text from PDF: " + repr(e))
    
def extract_texts_from_files(uploaded_files):
    """
    Given a list of uploaded files (e.g., for one fund),
    extract text from each file and combine them into a single string.
    """
    combined_text = ""
    for uploaded_file in uploaded_files:
        file_path = save_uploaded_file(uploaded_file, UPLOAD_FOLDER)
        # Here we assume the files are PDFs; you can extend this logic if needed.
        file_text = extract_text_from_pdf(file_path)
        combined_text += "\n" + file_text
    return combined_text.strip()    

# NEW: Process multiple funds for comparison against P1
def process_funds_for_comparison(funds_uploads, p1_files):
    if not p1_files:
        raise ValueError("No P1 files provided.")
    # Process the P1 files once (combine their text)
    p1_text = extract_texts_from_files(p1_files)
    comparison_results = []
    # Loop over each fund (each element is a list of uploaded files for that fund)
    for idx, fund_files in enumerate(funds_uploads):
        if fund_files:
            fund_text = extract_texts_from_files(fund_files)
            try:
                # Call your GPT function to compare this fund vs. P1
                comparison_result = extract_fund_comparison_with_gpt(
                    fund1_text=fund_text,
                    fund2_file1_text=p1_text,
                    fund2_file2_text=p1_text  # Passing P1 text twice as required
                )
            except Exception as e:
                comparison_result = f"Error processing Fund {idx+1}: {str(e)}"
            comparison_results.append((idx+1, comparison_result))
        else:
            comparison_results.append((idx+1, "No files uploaded for this fund."))
    return comparison_results
  

def extract_risk_details(file_path):
    """
    Extract risk level, type, first sentence, and last sentence from an uploaded image or document using OCR.
    
    Args:
    - file_path: Path to the uploaded image or document.
    
    Returns:
    - A dictionary with risk details (level, type, first sentence, last sentence).
    """
    try:
        # Use Tesseract OCR to extract text from the uploaded image
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)

        # Extract relevant details
        level_of_risk = re.search(r"Risk Level\s(\d+)", text)
        risk_level = re.search(r"Risk Type:\s([\w\s-]+)", text)

        # Extract first and last sentence from the "Definition of [Risk Type]" section
        definition_section = re.search(r"Definition of [\w\s-]+:(.+)", text, re.DOTALL)
        if definition_section:
            sentences = definition_section.group(1).strip().split(".")
            first_sentence = sentences[0].strip() + "." if len(sentences) > 0 else ""
            last_sentence = sentences[-1].strip() + "." if len(sentences) > 1 else ""
        else:
            first_sentence = ""
            last_sentence = ""
   
        return {
            "level_of_risk": level_of_risk.group(1) if level_of_risk else "Unknown",
            "risk_level": risk_level.group(1).strip() if risk_level else "Unknown",
            "first_sentence": first_sentence,
            "last_sentence": last_sentence,
        }
    except Exception as e:
        raise ValueError("Error extracting risk details: " + repr(e))
    
# 4) Main function that loops over multiple plan files
def process_plan_report(uploaded_files):
    """
    Loop over each uploaded file, extract text, call GPT, and gather plan data.
    If any file can't be parsed as docx/pdf/etc., log the file name and skip it.
    """
    all_plan_data = []

    for uf in uploaded_files:
        # 1) Save the file
        file_path = save_uploaded_file(uf, UPLOAD_FOLDER)

        # 2) Extract text, with try/except
        try:
            extracted_text = extract_text_from_file(file_path)
        except Exception as e:
            # Log or show an error in Streamlit, indicating which file failed
            st.error(f"Failed to extract text from '{uf.name}': {e}")
            # Optionally skip to next file
            continue

        # 3) Process extracted text with GPT
        try:
            plan_list = extract_plan_details_with_gpt(extracted_text)
            all_plan_data.extend(plan_list)
        except Exception as e:
            st.error(f"GPT error processing file '{uf.name}': {e}")
            # Optionally skip or continue
            continue

        # If successful:
        st.success(f"Successfully processed file: {uf.name}")

    return all_plan_data
    
def process_fund_reviews_single_prompt(uploaded_files):
    """
    For each file:
      - Save & extract text
      - Call generate_fund_review_single_call(extracted_text)
      - Collect each review in a list
    Returns a list of final text blocks (one per file).
    """
    all_reviews = []
    for uf in uploaded_files:
        file_path = save_uploaded_file(uf, UPLOAD_FOLDER)
        try:
            extracted_text = extract_text_from_file(file_path)
        except Exception as e:
            st.error(f"Failed to extract text from '{uf.name}': {e}")
            continue

        # Single GPT call that parses owner/fund + writes the final review
        try:
            review_text = generate_pension_review_section(extracted_text)
            all_reviews.append(review_text)
            st.success(f"Successfully generated review for {uf.name}")
        except Exception as e:
            st.error(f"GPT error on '{uf.name}': {e}")

    return all_reviews

def process_funds_for_comparison(funds_uploads, p1_files):
   
    if not p1_files:
        raise ValueError("No P1 files provided.")
    p1_text = extract_texts_from_files(p1_files)
    comparison_results = []
    for idx, fund_files in enumerate(funds_uploads):
        if fund_files:
            fund_text = extract_texts_from_files(fund_files)
            try:
                comparison_result = extract_fund_comparison_with_gpt(
                    fund1_text=fund_text,
                    fund2_file1_text=p1_text,
                    fund2_file2_text=p1_text  # Passing consolidated P1 text twice
                )
            except Exception as e:
                comparison_result = f"Error processing Fund {idx+1}: {str(e)}"
            comparison_results.append((idx+1, comparison_result))
        else:
            comparison_results.append((idx+1, "No files uploaded for this fund."))
    return comparison_results


def extract_client_details_with_gpt(factfinding_text):
    prompt = f"""
    You are an AI assistant tasked with extracting specific client details from a FactFinding report.

    **Objective**:
    Analyze the provided FactFinding report and extract the following details to populate placeholders in a financial document:

    **Placeholders**:
    - Full name: Combine Title and Surname and if the fact finds is adreesed for two people combine both names like : Mr forename+Surname & Mrs Forename+Surname ; example "Mr James Yeandle & Mrs Elizabeth Yeandle". 
    - Address: Full multiline address with postal code
    - Today's date: Current date in "9th January 2025" format
    - Salutation: "Dear [Forename]," format and if it is for two people use "Dear [Forename] & [Forename]," format.

    **FactFinding Report**:
    {factfinding_text}

    **Expected JSON Format**:
    {{
      "Full name": "",
      "Address": "",
      "Today’s date": "",
      "salutation": ""
    }}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        raw_content = response.choices[0].message.content
        cleaned_content = clean_json_response(raw_content)
        return parse_json_response(cleaned_content, "client details extraction")
    except Exception as e:
        error_msg = f"Client details error: {str(e)}"
        if 'raw_content' in locals():
            error_msg += f"\nRaw response: {raw_content}"
        raise ValueError(error_msg)
    

def generate_current_situation(factfinding_text):
    """
    Use OpenAI to generate a detailed 'Current Situation' section based on the FactFinding report.
    """
    prompt = f"""
    You are an assistant tasked with creating a "Current Situation" summary for a financial report.

    **Objective**:
    1. Analyze the provided FactFinding report and extract all relevant details.
    2. Write a detailed "Current Situation" section summarizing the client's financial and personal details.
    3. Write the "Current Situation" section as bullet points only. Do not include any headings, introductions.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
   - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
   - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
   - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".    
    - If the report is for one individual, use singular language (e.g., "Chris, you are 68 years....","You have...", "Your pension is...").
    - If the report is for two individuals, use plural or joint language (e.g., "Tony, you are 74 years old, and Liz, you are 75 years old.... ", "You both have...", "Your combined pensions are...").       
    - Write the section in professional bullet points, but keep it conversational by using "You" at the beginning of some sentences.
    - Make sure to mention the part about include monthly growth and monthly expenditure and remain in one line for example (• You have a monthly gross income of £2,700.00 and a monthly expenditure of £1,670.00, leaving you with a monthly surplus of £1,030.00.).
    - Include point about dependants and Wife details and their financial details if they are found.
    - When extracting dependents' details, exclude the name of the person who took the notes (e.g., "16 Sep 2024 - Alex Armstrong"). Focus only on the actual dependents' names and their details.
    - Whenever you mention the phrase "You should always have 3-6 months worth of expenditure," always calculate and include the range for 3-6 months of monthly expenditure, formatted as (£<3 months value>-£<6 months value>). Use the extracted monthly expenditure for this calculation(3 multiple by monthly expenditure - 6 multiple by monthly expenditure ).
    - Include line about Protection detalis (e.g,  Home insurance, car insurance,..... .)
    - Include sufficient length and detail, following this example:

    **Example "Current Situation" Section**:
    
    Current Situation

    When we last met, you made me aware of your current situation:

    • Chris, you are 68 years old, co-habiting with your partner and in good health. 
    • You retired in January 2020 having previously worked in IT.  
    • You own your house outright which is worth approximately £555,000.00.
    • You are in receipt of your full state pension (£930.00 gross per month), Zurich Financial Services final salary pension (£640.00 gross per month) and your NHS pension (£890.00 gross per month).
    • You withdraw £800.00 gross per month from your tax-free cash entitlement from your Royal London Personal Pension. This will be exhausted soon. 
    • You have no debts or other liabilities.
    • You have a monthly gross income of £3,260.00 and a monthly expenditure of £2,410.00, leaving you with a monthly disposable of £850.00. 
    • You have no financial dependents. 
    • You have £39,000.00 in cash reserves. This is a sufficient emergency fund. You should always have 3-6 months worth of expenditure in an easy access bank account for emergencies (£7,230.00-£14,460.00). You may want to invest any over this amount to get better returns than cash in the bank. 
    • You have drafted a Will and Power of Attorney, and they are both up to date.

    **Your Task**:
    1. Ensure your output matches the tone, structure, and level of detail of the example provided above.
    2. Extract details from the FactFinding report, including:
       - Personal details (age, marital status, and health).
       - Dependants and Wife details and their financial details if they are found.
       - Include point about dependants and Wife details and their financial details if they are found.
       - Retirement details and previous occupation , also for his wife it was found in one line with his detalis.
       - Property ownership and approximate value.
       - Pensions, incomes, and cash withdrawals.
       - Include point about monthly growth and monthly expenditure in one line, you can use the Incomes and Expenses table (The Sum of each), use the word growth and exenditure when talking about it.
       - Emergency funds and recommendations for improvement.
       - Succession planning (Will and Power of Attorney).
       - Protection
       - Any other relevant details.

    **FactFinding Report**:
    {factfinding_text}

    **Output**:
    Write a "Current Situation" section that:
    1. Matches the tone and format of the example.
    2. Includes sufficient detail to make the section comprehensive and professional.
    3. Avoids mentioning specific details about thier invstment knowledge and experience.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return response.choices[0].message.content


def generate_priorities_and_objectives(factfinding_text):
    """
    Use OpenAI to generate a detailed 'Priorities and Objectives' section based on the FactFinding report.
    """
    prompt = f"""
    You are an assistant tasked with creating a "Priorities and Objectives" section for a financial report.

    **Objective**:
    1. Analyze the provided FactFinding report and extract relevant details fro m the Objectives table.
    2. Personalize the "Priorities and Objectives" section by integrating other relevant information in the report, such as financial circumstances, retirement plans, health, and family situation.
    3. Write the "Priorities and Objectives" section as a bullet points or numbers for example (e.g., 1. lab akbdh sj h , 2. hushfjfhsuhfivn;k , 3.h osijck,........). Do not include headings or numbered lists.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
    - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
    - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
    - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".   
    - If the report is for one individual, use singular language (e.g., "You have...", "Your pension is...").
    - If the report is for two individuals, use plural or joint language (e.g., "You both have...", "Your combined pensions are...").   
    - Focus on the client's primary and secondary financial objectives, integrating personal context to make it highly personalized.
    - Use a professional tone and structure, with clear and specific details.
    - Avoid unnecessary repetition or vague language.
    - Ensure that the content remains cohesive and logical.
    - Write the "Priorities and Objectives" section as a bullet points or numbers for example (e.g., 1. lab akbdh sj h , 2. hushfjfhsuhfivn;k , 3.h osijck,........).


    **Your Task**:
    1. Ensure your output matches the tone, structure, and level of detail of the example provided above.
    2. Extract details from the FactFinding report, including:
       - Financial objectives (e.g., maintaining living standards, creating income, building capital).
       - Any specific income goals or plans (e.g., covering monthly expenses or specific purchases).
       - Retirement plans or strategies (e.g., utilizing pensions, maintaining capital for long-term care).
       - Any family or personal considerations that impact objectives (e.g., no descendants, focus on enjoying funds in the client's lifetime).
       - Relevant details from the Objectives table, such as priority goals, timeframes, and preferences.
    3. Write the section in a professional and concise manner.

    **FactFinding Report**:
    {factfinding_text}

    **Output**:
    Write a "Priorities and Objectives" section that:
    1. Matches the tone and structure of the example provided above.
    2. Includes sufficient personalization based on the Objectives table and other relevant details.
    3. Avoids unnecessary repetition or vague statements.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return response.choices[0].message.content


def generate_multi_risk_attitude_text(extracted_texts):

    if not extracted_texts:
        return "No risk profile text provided."

    text_blocks_json = json.dumps(extracted_texts, indent=2)

    prompt = f"""
You are an AI that processes multiple risk profile text blocks and creates a single 'Attitude to Risk' statement.

First, parse the following list of text blocks (in JSON). Each text block is from a separate risk profile:
{text_blocks_json}

For each text block i upload, you must extract:
  1. First Name
  2. Risk Level (1-5)
  3. Risk Type (e.g., 'Cautious' or 'Balanced')
  4. First sentence of risk definition
  5. Last sentence of risk definition

After extracting these details for each block, determine if ALL individuals share the SAME risk level (Scenario A) or NOT (Scenario B).

Scenario A (all same risk level):
Use this exact text (substituting the actual risk level/type):
"The risk profiler completed with you both came out with an attitude to risk level {{risk_level}} which is {{risk_type}}.
As {{risk_type}} investors, you do not see yourself as a particularly cautious person and have no strong or negative
associations with the notion of taking risk. You can be inclined to look for a combination of investments with differing
levels of risk and understand that you may need to take some risk to meet your investment goals."

Scenario B (different risk levels):
For each person individually, produce:
"{{first_name}}, the risk profiler completed with you came out with an attitude to risk level {{risk_level}} which is {{risk_type}}.
{{first_sentence}}. {{last_sentence}}."

If there is only one text block, you can choose either scenario A or B. (e.g., treat a single user like scenario B.)

Finally, return ONLY the final combined text. 
    """

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        final_text = response.choices[0].message.content.strip()
        return final_text
    except Exception as e:
        raise ValueError(f"Error generating multi-risk text: {e}")
      

# 3) GPT prompt for plan details
def extract_plan_details_with_gpt(extracted_text):
    """
    Sends the extracted text to GPT, asking for plan details in JSON array form:
    [
      {
        "Provider": "some provider",
        "Plan Number": "123456",
        "Plan Type": "personal Pension",
        "Current Value": "£100,000"
      },
      ...
    ]
    """
    prompt = f"""
Extract plan details from the text below and return a JSON array of objects. Example structure:
[
  {{
    "Provider": "XYZ",
    "Plan Number": "12345",
    "Plan Type": "Personal Pension",
    "Current Value": "£210,000"
  }}
]

Text:
{extracted_text}
"""
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    raw_content = response.choices[0].message.content
    cleaned = clean_json_response(raw_content)

    # Attempt to parse as JSON
    try:
        return json.loads(cleaned)
    except:
        # If parsing fails or GPT didn't return valid JSON, return empty
        return []
       
    
def generate_pension_review_section(extracted_text):
    """
    Generates a comprehensive 'Review of Existing Fund' section per plan heading,
    ensuring no duplicate reviews for the same plan number, ignoring overall
    portfolio summaries, and properly handling ISAs vs. Pensions/Bonds.
    """

    prompt = f"""
You are a financial assistant tasked with generating a professional "Review of Existing Fund" section for each plan heading found in the text, **ignoring** any overall or combined summary.

---
### **Key Requirements**:
1. **Identify each plan heading** (e.g., "P0699-30007-00787 - ISA") in the text. 
   - Extract only the **last 5 digits** of the plan number for display. 
   - Example: "P0699-30007-00787" → "00787".
2. **Ignore** any overall portfolio summaries or total amounts that lump all plans together. 
   - We only want plan-specific details under each heading.
3. **Never output more than one review** for the same plan number. If a plan heading repeats, produce only one review.
4. **Extract** details from the table below each plan heading:
   - Start Date (if available)
   - Current Value
   - Total Investment
   - Investment In / Out
   - Change (£ and %)
   - Recent Change (£ and %) if shown
   - Term (if shown)
   - Attitude or risk level (if shown)
   - Withdrawals (if applicable)
5. **If withdrawals exceed the 5% allowance (for Bonds/Pensions)**, include:
   > "By taking withdrawals over the 5% allowance, you trigger a chargeable event,
   leaving you open to a potential income tax liability. As long as your total
   income remains in the basic rate income tax band, you should not have to pay
   any additional income tax, as this has already been considered to have been
   paid by [Provider Name]."
6. **If 'lives assured on death' is mentioned**, include:
   > "[People covered by it] are lives assured on the bond, meaning that the bond
   will remain in force until the second of you passes away, regardless of who
   owns the bond or makes withdrawals. This is therefore not a capital redemption
   bond as discussed in our meeting."
7. **Use correct language** for each fund type:
   - **ISA** → emphasise tax efficiency, not pension withdrawals.
   - **Bond/Pension** → mention possible chargeable events, retirement, etc.
8. **British English** grammar/spelling; no bullet points (full paragraphs). 
9. **Do not merge** multiple funds into one review. Each heading → one distinct review.
10. **Never copy the example text** directly; only use it as a format guide.

---
### **Example Format**:
**Review of Existing [Fund Type] - [Owner First Name]**

Your [Fund Type] under Plan Number **[Last 5 Digits]** was opened on **[Start Date]** and is currently valued at **£[Current Value]**. This reflects an overall change of **£[Change]** (**[Change %]**) since your initial investment of **£[Investment In]**.

You have contributed a total of **£[Total Investment]** into this fund over the term. [Any mention of withdrawals, tax, 5% allowance, or lives assured clause if relevant.]

---
### **Your Task**:
1. **Parse** the input text to find each plan heading + table.
2. **Generate exactly one** "Review of Existing [Fund Type]" per plan heading. 
3. **Ignore** any overall or combined totals that summarize all plans.
4. **Output** full paragraphs for each plan, referencing the correct plan number (last 5 digits).
5. **Don’t** produce duplicates for repeated headings.
6. Insert the 5% allowance statement or lives assured statement if the data indicates.

---
**Product Report**:
{extracted_text}

---
**Output**:
- A separate 'Review of Existing [Fund Type] - [Owner Name]' section **for each unique plan number**.
- Each section must reflect the plan’s table details (start date, value, etc.).
- Maintain a clear, concise, professional tone and follow the instructions above.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )

    return response.choices[0].message.content.strip()


def extract_investment_portfolio_with_gpt(extracted_text):
    """
    Use GPT to extract the portfolio's total investment and the table of individual fund holdings.
    The prompt instructs GPT to:
      1. Locate the section of the text that contains the overall portfolio report. Typically, this section starts with a heading such as "Portfolio Report (date)" followed by a line indicating "Total Investment" and its amount.
      2. Extract the portfolio's total investment amount.
      3. Locate the table of individual fund details. The table headers might include various terms such as "Fund Description", "Total Investment", "Total Value", "Asset Value", "Units", "Unit Price", "Change", or "Recent Change". These headers may vary but convey similar meanings.
      4. For each fund, extract at least the fund name (from a header like "Fund Description") and the fund’s value (from "Total Value", "Asset Value", or a similar header).
      5. Calculate the percentage that each fund’s value represents out of the total investment.
      6. Append a final "TOTAL" row that shows 100%.
      
    The expected output is a JSON object in the following format:

    {
      "PortfolioTotal": <number>,
      "Holdings": [
         { "Fund": "<Fund Name>", "Value": <number>, "Percent": "<percentage string>" },
         ...,
         { "Fund": "TOTAL", "Value": <number>, "Percent": "100%" }
      ]
    }

    Do not include any additional text or markdown formatting.

    Text to analyze:
    {extracted_text}
    """
    prompt = f"""
You are an AI assistant tasked with analyzing investment portfolio data provided as plain text.
Follow these steps:
1. Locate the section in the text that contains the portfolio summary. This section usually starts with a heading like "Portfolio Report" (which may include a date) and then a line such as "Total Investment" followed by an amount (e.g., "£123,456.78").
2. Extract the portfolio's total investment amount.
3. Next, find the table (or section) that lists individual fund details. The table headers might vary but will include labels such as "Fund Description", "Total Investment", "Total Value", "Asset Value", "Units", "Unit Price", "Change", or "Recent Change". Even if these labels are not exactly the same, they convey the same meaning.
4. For each fund listed, extract at least:
   - The fund name (from the column that describes the fund, e.g., "Fund Description")
   - The fund's value (from the column that indicates "Total Value" or "Asset Value")
5. Calculate each fund’s percentage of the total investment (fund value divided by total investment), and format it as a percentage string.
6. Add a final row labelled "TOTAL" with the total investment amount and "100%".

Return a JSON object exactly in the following structure:

{{
  "PortfolioTotal": <number>,
  "Holdings": [
    {{
      "Fund": "<Fund Name>",
      "Value": <number>,
      "Percent": "<percentage string>"
    }},
    ...,
    {{
      "Fund": "TOTAL",
      "Value": <number>,
      "Percent": "100%"
    }}
  ]
}}

Do not include any additional text or markdown formatting.

Text to analyze:
{extracted_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        raw_content = response.choices[0].message.content
        # Remove any markdown code fences if present
        cleaned = re.sub(r'^```json\s*|\s*```$', '', raw_content, flags=re.DOTALL).strip()
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse investment portfolio GPT response: {str(e)}\nRaw response: {raw_content}")
    except Exception as e:
        raise ValueError(f"Investment portfolio extraction error: {str(e)}")

def generate_safe_withdrawal_rate_sections(plan_texts: list) -> list:
    """
    Given a list of plan report texts (one per file), generate the Safe Withdrawal Rate section
    for each one by calling the existing generate_safe_withdrawal_rate_section.
    Returns a list of SWR section strings.
    """
    swr_sections = []
    for text in plan_texts:
        # Call the existing single-text function for each file
        swr = generate_safe_withdrawal_rate_section(text)
        swr_sections.append(swr)
    return swr_sections


def generate_safe_withdrawal_rate_section(plan_report_text):
    """
    Use OpenAI to generate a detailed 'Safe Withdrawal Rate (SWR)' section
    based on the plan report.

    Args:
        plan_report_text (str): Extracted text from the plan report.

    Returns:
        str: The generated SWR section or an empty string if no withdrawals are detected.
    """
    # Ensure your OpenAI API key is set in the environment variables
    
    prompt = f"""
    You are a financial assistant tasked with creating a "Safe Withdrawal Rate (SWR)" section for a financial report.

    **Objective**:
    1. Analyze the provided plan report text to determine if the client is withdrawing money from their investments (e.g., pensions, savings, other investments).
    2. If withdrawals are present, extract the following details:
       - Monthly withdrawal amount.
       - Annual withdrawal rate.
       - Total portfolio value.
    3. Generate a detailed and personalized "Safe Withdrawal Rate (SWR)" section based on the extracted information.
    4. Use a conversational tone that feels direct and engaging, starting sentences with "You" where appropriate.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
    - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
    - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
    - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".    
    - If withdrawals are detected:
      - Calculate the current withdrawal rate as (annual withdrawals / total portfolio value) * 100, dont write the equation just the results.
      - Compare the current withdrawal rate to the generally accepted safe withdrawal rate of 4.00%.
      - Highlight the risks associated with withdrawing more than the recommended rate.
    - If no withdrawals are detected, respond with "No withdrawals detected."

    - Write clearly and concisely in full paragraphs, avoiding bullet points.
    
- **Important**: Only consider **monthly** withdrawals when calculating the SWR. If the text only mentions an annual or lump-sum withdrawal, treat it as if there is no monthly withdrawal.
  - If a monthly withdrawal is detected:
    - Calculate the current withdrawal rate as (12 × monthly withdrawal) / (total portfolio value) × 100 (do not show the formula; just show the resulting percentage).
    - Compare this monthly-based annual rate to the generally accepted safe withdrawal rate of 4.00%.
    - Highlight the risks associated with withdrawing more than the recommended rate.
  - If there is no mention of a monthly withdrawal (or only mention of annual or lump-sum withdrawals), respond with: "No monthly withdrawals detected."


    **Example Section**:
    Safe Withdrawal Rate (SWR)

    You're taking a taxable income of £1,200.00 from your plan every month.

    The generally accepted safe withdrawal rate for retirement income is approximately 4.00% per year of the portfolio’s value.

    Your current withdrawal rate is: 11.61%.

    Withdrawing more than 4.00% annually may deplete your investment faster than anticipated, especially during periods of market volatility.

    With a withdrawal rate exceeding 4.00%, there is an increased risk that your investment may not last throughout your retirement. This could result in a shortfall in later years, reducing your ability to meet essential expenses. We recommend regular reviews to ensure your withdrawals remain sustainable.

    **Plan Report Text**:
    {plan_report_text}

    **Output**:
    Write a "Safe Withdrawal Rate (SWR)" section that:
    - Follows the structure and tone of the example provided.
    - Includes all relevant extracted details.
    - Provides a clear comparison between the client's withdrawal rate and the recommended rate.
    """

    try:
        # Make a call to the OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a financial advisor assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,  # Lower temperature for more deterministic output
            max_tokens=600,    # Adjust as needed to capture detailed responses
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        # Extract the response text
        generated_text = response.choices[0].message.content.strip()

        if generated_text == "No withdrawals detected.":
            return ""  # No SWR section needed
        else:
            return generated_text

    except Exception as e:
        print("Error generating SWR section: " + repr(e))
        return ""
    
    
def process_single_fund_performance(text: str):
    """
    Process a single fund performance text and return the parsed JSON data.
    """
    prompt = f"""
Analyze this fund performance data and return JSON with:
1. Yearly performance percentages
2. Benchmark comparisons
3. Cumulative 5-year sum (calculated as simple sum of yearly percentages)

Rules:
- Use EXACTLY this format:
[
    {{
        "Fund": "Fund Name",
        "Year 1": "X%",
        "Year 2": "X%",
        "Year 3": "X%",
        "Year 4": "X%",
        "Year 5": "X%",
        "Cumulative (5 YR)": "X%",
        "Benchmark": {{
            "Year 1": "X%",
            "Year 2": "X%",
            "Year 3": "X%",
            "Year 4": "X%",
            "Year 5": "X%",
            "Cumulative (5 YR)": "X%"
        }}
    }}
]
- Only return raw JSON without any additional text or markdown.
- Use "N/A" for missing data.

Text to analyze:
{text}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    raw_content = response.choices[0].message.content
    # Remove any markdown code fences if present
    cleaned_content = raw_content.strip().replace('```json', '').replace('```', '')
    return json.loads(cleaned_content)

def extract_fund_performance_with_gpt(extracted_texts):
    """
    Accepts either a single string (extracted text from one file) or a list of strings (one per file).
    Processes each text individually using process_single_fund_performance.
    
    Returns:
    - A list of JSON objects if multiple texts are provided.
    - A single JSON object if only one text is provided.
    """
    # Multi-file case: if extracted_texts is a list
    if isinstance(extracted_texts, list):
        results = []
        for text in extracted_texts:
            if text.strip():
                results.append(process_single_fund_performance(text))
        return results
    else:
        # Single file case
        if extracted_texts.strip():
            return process_single_fund_performance(extracted_texts)
        else:
            return None

def process_single_dark_star_performance(text: str):
    """
    Process a single Dark Star performance text and return the parsed JSON data.
    """
    prompt = f"""
    You are an AI assistant tasked with extracting fund performance details from a financial report.
    Analyze the text below and return a JSON response similar to this format:
    [
        {{
            "Fund": "Dark Star Asset Management Balanced Plus",
            "Year 1": "15%",
            "Year 2": "9.7%",
            "Year 3": "6.5%",
            "Year 4": "4.3%",
            "Year 5": "N/A",
            "Benchmark": {{}},
            "Cumulative (5 YR)": "33.9%"
        }}
    ]

    If no benchmark is provided in the text, leave the "Benchmark" field empty.
    **Important Instructions**:
        - Use the "YTD" (Year-to-Date) data for yearly performance.
        - Map the most recent year as Year 1, the next as Year 2, and so on.
        - If a year has no data, use "N/A".
        - Cumulative performance should only include available years and should be calculated as the sum of those percentages.

    Text:
    {text}
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw_content = response.choices[0].message.content
        # Clean out any markdown code fences if present
        cleaned_content = re.sub(r'^```json\s*|\s*```$', '', raw_content, flags=re.DOTALL)
        return json.loads(cleaned_content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Dark Star JSON error: {e}\nResponse: {raw_content}")

def extract_dark_star_performance_with_gpt(extracted_texts):
    """
    Accepts either a single string (extracted text from one file) or a list of strings (each from one file).
    Processes each text individually using process_single_dark_star_performance.
    
    Returns:
    - A list of JSON objects if multiple texts are provided.
    - A single JSON object if only one text is provided.
    """
    # Multi-file case
    if isinstance(extracted_texts, list):
        results = []
        for text in extracted_texts:
            if text.strip():
                results.append(process_single_dark_star_performance(text))
        return results
    else:
        # Single file case
        if extracted_texts.strip():
            return process_single_dark_star_performance(extracted_texts)
        else:
            return None

    

def extract_sap_comparison_with_gpt(extracted_text):
    """
    Extract comparison details, age, and unify row labels using a detected company name.
    The prompt instructs GPT to search for headings like 'Accumulative Comparison' or 'Comparison at Age X'
    and to normalize row labels for key metrics using the detected company name.
    """
    prompt = f"""
    You are an AI assistant that extracts a single comparison table and the relevant age from a financial SAP report.

    The text may have headings such as:
      - 'Accumulative Comparison'
      - 'Accumulative Comparison at Age X'
      - 'Comparison at Age X'
      - 'Projected Fund (Paid Up) at proposed age X'

    Steps:
    1. Identify the heading that says 'Accumulative Comparison' or 'Comparison at Age X' or includes an age.
    2. Extract ONLY the table immediately below that heading (until a blank line or a new heading).
    3. Find the projected age from the heading (e.g. "Comparison at Age 80" means Age=80).
    4. Identify the company name from row labels (for example, if you see "Royal London Pension Portfolio – Profit Share", 
       or "Rate of Return Required from Quilter", then detect the company name as "Royal London" or "Quilter").
    5. Normalize the row labels by replacing the company name in the key metrics with the detected company name. 
       For example, if the company is "Royal London", the rows should be:
         - "Rate of Return Required from Royal London"
         - "Effect on Fund if Moved to Royal London"
         - "Reduction in Yield if Moved to Royal London"
         - and keep any other rows as is.
    6. Ignore any other tables in the text.

    Return a JSON object in this structure:
    {{
      "Age": 80,
      "companyName": "Royal London",
      "Table": {{
         "Assumed Growth Rates": ["2%", "5%", "8%"],
         "Existing Schemes": ["£118,972.00", "£155,558.00", "£201,866.00"],
         "Rate of Return Required from Royal London": [...],
         "Effect on Fund if Moved to Royal London": [...],
         "Reduction in Yield if Moved to Royal London": [...],
         "Royal London Pension Portfolio - Profit Share": [...]
      }}
    }}

    If a specific row is not found, keep the label as is. Use 'Unknown' if you cannot determine the company name.

    Text to parse:
    {extracted_text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw_content = response.choices[0].message.content
        # Clean any markdown fences
        cleaned_content = re.sub(r'^```json\s*|\s*```$', '', raw_content, flags=re.DOTALL)
        return json.loads(cleaned_content)
    except json.JSONDecodeError as e:
        raise ValueError(f"SAP comparison JSON error: {e}\nResponse: {raw_content}")

def extract_annuity_quotes_with_gpt(extracted_text):
    """
    Use GPT to extract annuity quotes details from the extracted text and return a plain text response.
    """
    prompt = f"""
        ### Instructions:
        Extract the details of annuity quotes from the provided text and format the output as follows:

        Example Output:
        Quote 1:
        - Purchase Amount: £124,030
        - Monthly Amount: £854
        - Yearly Amount: £10,250
        - Yearly Increase: None

        Quote 2:
        - Purchase Amount: £124,030
        - Monthly Amount: £603
        - Yearly Amount: £7,242
        - Yearly Increase: Retail Price Index (RPI)

        Quote 3:
        - Purchase Amount: £124,030
        - Monthly Amount: £668
        - Yearly Amount: £8,018
        - Yearly Increase: 3.00%

        ### Notes:
        1. **Purchase Amount**:
        - Locate the number next to the text `pension pot`.
        - For example, if you see `£124,030 pension pot`, extract `£124,030`.

        2. **Monthly Amount**:
        - Find the number followed by the word `monthly` under the **Your Income** section.
        - For example, if you see `£854 monthly`, extract `£854`.

        3. **Yearly Amount**:
        - Find the number followed by the word `yearly` under the **Your Income** section.
        - For example, if you see `£10,250 yearly`, extract `£10,250`.

        4. **Yearly Increase**:
        - Look for the yearly increase information in the **Your Choices** section.
        - Extract as:
            - `"None"` for `No annual increase`.
            - `"Retail Price Index (RPI)"` for `Increase by RPI`.
            - `"3.00%"` for `Increase 3% per year`.

        ---

        Text to Analyze:
    {extracted_text}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    # Return the formatted plain text output
    return response.choices[0].message.content


def extract_fund_comparison_with_gpt(fund1_text, fund2_file1_text, fund2_file2_text):
    """
    Generate a GPT response for fund comparison in the specified short format.
    """
    prompt = f"""
    ### Role
    You are a financial data extraction system specialized in processing multi-file texts from investment fund reports. Your task is to analyze the provided text for Royal London and P1 funds and extract fee metrics from sections that start with terms like "Fees and charges", "Service Charges", or "COSTS AND CHARGES ANNUAL SUMMAR". 

    ### Metrics to Extract:
    1. **Plan Value**: Locate the plan value associated with the review period by looking for text after "Review dates:".
    2. **Weighted Fund Charge (WFC)**:
       - For Royal London: Look for phrases like "equivalent to X% of the value of your plan each year." Extract X%.
       - For P1: If not mentioned, default to 0.44%.
    3. **Platform Charge**:
       - Search for any instance of the word "Platform" or phrases like "[Provider Name] charges" (e.g., "Quilter charges").
       - Extract the percentage value if available; otherwise, default to 0.0%.
    4. **Ongoing Advice Fee (OAF / Advice Charges)**:
       - Look for terms such as "Ongoing Advice Fee", "OAF", "advice fee", or "Advice charges".
       - Always use a default value of 0.50% if no specific value is found.
    5. **Discretionary Management Charges**:
       - Look for any fee that includes the word "Discretionary" (such as "Discretionary Manager fees" or "Discretionary management charges").
       - For Royal London, if not mentioned, default to 0.0%.
    6. **Drawdown Fee**:
       - Search for terms like "Drawdown" or "Product" to extract the fee.
    7. **ProfitShare**:
       - For Royal London, always set ProfitShare to -0.15%.
       - For P1, search for its value; if not mentioned, default to 0.0%.

    ### Input Text:
    **Royal London Details**:
    {fund1_text}

    **P1 Details**:
    {fund2_file1_text}
    {fund2_file2_text}

    ### Tasks:
    - First, identify the section in the text that starts with "Fees and charges", "Service Charges", or "COSTS AND CHARGES ANNUAL SUMMAR".
    - Within that section, extract each of the metrics listed above along with their percentage values.
    - If a metric is not mentioned in the text, use the default value as specified.
    - Multiply each percentage by the Royal London Plan Value to calculate monetary amounts.
    - Sum these amounts to determine the Total Annual Ongoing Charges for each fund.
    - Compute the differences between the P1 fund and the Royal London fund for each fee category and overall totals.
    - Generate a dynamic comparison statement summarizing the fee differences.

    ### Output Format:
    Return **ONLY** the populated template below with the extracted and calculated values. Do not include any additional explanations or markdown formatting:

    ---
    Plan value = £[Value]  
    Weighted Fund Charge % = [Value]% (£[Value])  
    Platform Charge % = [Value]% (£[Value])  
    Ongoing Advice Fee % = 0.50% (£[Value])  
    Discretionary Fund Manager Charge % = [Value]% (£[Value])  
    Drawdown Fee % = [Value]% (£[Value])  
    ProfitShare % = -0.15% (£[Value])  

    **P1 Metrics**:  
    Weighted Fund Charge % = [Value]% (£[Value])  
    Platform Charge % = [Value]% (£[Value])  
    Ongoing Advice Fee % = [Value]% (£[Value])  
    Discretionary Fund Manager Charge % = [Value]% (£[Value])  
    Drawdown Fee % = [Value]% (£[Value])  
    ProfitShare % = [Value]% (£[Value])  

    **Total Annual Ongoing Charges**:  
    - Royal London: [Total_%]% (£[Total_£])  
    - P1: [Total_P1_%]% (£[Total_P1_£])  

    **Comparison**: [Dynamic_Statement]  
    ---

    **Rules**:
    - Use "0.0%" for any missing values.
    - Format all monetary values as £1,234.56.
    - Do not include any additional commentary or explanations.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    return response.choices[0].message.content

        
                                
def extract_last_year_performance_text(extracted_text: str) -> str:
    """
    Searches the text for a 'Year on year performance' section and tries to parse
    the last year’s growth percentage. Returns a short sentence summarising the last year.

    Example return:
        "Your fund returned 10.2% in the last year."
    or
        "Unable to find last year performance in the text."

    """
    prompt = f"""
    You are an AI assistant that extracts the most recent 'year on year performance'
    from the provided text.

    The text often includes a table or lines like:

        Year on year performance
         30/09/2023-30/09/2022   30/09/2022-30/09/2021   30/09/2021-30/09/2020
         Fund growth%           10.2                   6.0                 -0.4

    Or it might read in a single line: "Fund growth%  10.2%  6.0%  ..."

    Your task:
    1. Identify the most recent one-year growth figure from the text (e.g. 10.2%).
    2. Return a short sentence in plain text: "Your fund returned 10.2% in the last year."

    If you cannot find a single last-year figure, respond with:
    "Unable to find last year performance in the text."

    Text to parse:
    {extracted_text}
    """

    try:

        # Call GPT with your prompt
        response = client.chat.completions.create(
            model="gpt-4o-mini",  
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        final_text = response.choices[0].message.content.strip()
        return final_text

    except Exception as e:
        # In case GPT fails or the parse fails
        return f"Unable to find last year performance in the text. Error: {e}"



def generate_iht_section(factfinding_text, plan_texts: list) -> str:
    """
    Given a FactFinding text and a list of plan report texts, 
    concatenate all plan texts and generate a single IHT section.
    """
    # Concatenate all plan texts into one string (separated by newlines)
    combined_plan_text = "\n".join(plan_texts)
    
    prompt = f"""
Analyze the provided data and calculate the IHT liabilities, taking into account the following UK IHT rules and thresholds. Use the client’s address to estimate property value based on typical property prices for the area,
search the `factfinding_text` for details about the client’s wife, dependents, mortgage, and debts, and include these details in the calculations. and the `combined_plan_text` for investment details.

**Objective**:
1. Analyze the provided FactFinding document and the combined plan report text.
2. Nil Rate Band (NRB)**:
   - Every individual has a Nil Rate Band (NRB) allowance of £325,000. This amount is **not taxed**.
   - For married couples or civil partners, the unused portion of the NRB can be transferred to the surviving spouse, effectively doubling the NRB to £650,000.

3.Residence Nil Rate Band (RNRB):
   - If the estate includes a main residence passed to direct descendants (e.g., children or grandchildren), an additional **Residence Nil Rate Band (RNRB)** of up to £175,000 may apply.
   - The RNRB is transferable between spouses, effectively doubling to £350,000 for married couples or civil partners.
   - If there is **no property**, the RNRB is not applicable and should be set to £0.00.

4. Note that the combined plan report may include details for multiple investments and, if applicable, details for two individuals (e.g. husband and wife). You can determine this by checking if names like "James" and "Elizabeth" appear in the texts.
5. For every investment found, output one line in the format:
   - If two names are present: "[Name] [Investment Type] ([Plan Number]) = £[Current Value]"
   - If only one person is present, simply: "[Investment Type] ([Plan Number]) = £[Current Value]"
6. Wife and Dependents:
   - Search the `factfinding_text` for information about the client's wife or civil partner and dependents.
   - If a wife or civil partner is found, include their NRB (£325,000) and RNRB (£175,000) in the calculations.
   - Use details about dependents (e.g., children or grandchildren) to determine eligibility for the RNRB.

7.Mortgage and Debts:
   - Search the `factfinding_text` for any mortgage, loans, or debts.
8. Pensions:
   - Include pensions invesments values where applicable from the product report.     
9. Then calculate the summary values:
   - Total Taxable Estate = Sum of all current values.
   - Mortgage and Debts: Extract from the FactFinding document (assume £0.00 if none).
   - If one person use Nil Rate Band if more than one (husbund, wife, childern) use Nil Rate Band x2 = £<value>
   - If one person use Residence Nil Rate Band if more than one (husbund, wife, childern) use Residence Nil Rate Band x2 = £<value>
   - Remaining Estate = Total Taxable Estate - (Nil Rate Band x2 + Residence Nil Rate Band x2 + Mortgage and Debts). If negative, set to £0.00.
   - Deduct Mortgage and Debts amounts from the total taxable estate to calculate the remaining estate.
   - Tax @ 40% = 40% of the Remaining Estate.
10. Output the results as bullet points, one per line.

**Input Data**:
1. FactFinding Document:
{factfinding_text}

2. Combined Plan Report Text:
{combined_plan_text}

**Output Format**:
Return only bullet points with each investment line and the summary rows exactly in this format (do not include any extra commentary):
- [Main Residence worth]
- [Mortgage and Debts]
- [Investment Line 1]
- [Investment Line 2]
...
- Total Taxable Estate = £<value>
- Mortgage and Debts = £<value>
- Nil Rate Band = £<value>
- Residence Nil Rate Band = £<value>
- Remaining Estate = £<value>
- Tax @ 40% = £<value>
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        iht_text = response.choices[0].message.content.strip()
        return iht_text
    except Exception as e:
        raise RuntimeError("Error generating IHT section: " + repr(e))


    
def create_plan_report_table(doc, plan_report_data):
    """
    Insert a table into the doc for the given plan data.
    plan_report_data is expected to be a list of dicts, each with
    { "Provider": "", "Plan Number": "", "Plan Type": "", "Current Value": "" }.
    """
    doc.add_heading("Plan Report Details", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = "Provider"
    hdr_cells[1].text = "Plan Number"
    hdr_cells[2].text = "Plan Type"
    hdr_cells[3].text = "Current Value"

    for plan in plan_report_data:
        row_cells = table.add_row().cells
        row_cells[0].text = plan.get("Provider", "")
        row_cells[1].text = plan.get("Plan Number", "")
        row_cells[2].text = plan.get("Plan Type", "")
        row_cells[3].text = plan.get("Current Value", "")

def create_comparison_table(document, sap_comparison_dict):
    """
    Generate the 'Comparison at Age {age}' table in the Word document.
    Expects sap_comparison_dict to have a key "Table" with the comparison data.
    It uses all keys (except "Assumed Growth Rates") as row labels.
    """
    if not sap_comparison_dict or "Table" not in sap_comparison_dict:
        raise ValueError("SAP comparison table data is missing or incorrectly formatted.")

    table_data = sap_comparison_dict["Table"]
    growth_rates = table_data.get("Assumed Growth Rates", [])
    if not growth_rates:
        raise ValueError("No 'Assumed Growth Rates' row found or it is empty.")

    # Use all keys except "Assumed Growth Rates" as rows
    row_labels = [key for key in table_data.keys() if key != "Assumed Growth Rates"]
    header = [""] + growth_rates

    table = document.add_table(rows=len(row_labels) + 1, cols=len(header))
    table.style = "Table Grid"

    # Fill in the header row
    for col_idx, hdr in enumerate(header):
        table.cell(0, col_idx).text = hdr

    # Fill in each row
    for row_idx, label in enumerate(row_labels, start=1):
        table.cell(row_idx, 0).text = label
        row_data = table_data.get(label, [])
        for col_idx in range(len(growth_rates)):
            val = row_data[col_idx] if col_idx < len(row_data) else "N/A"
            table.cell(row_idx, col_idx + 1).text = val

    return table



def add_investment_holdings_tables(document: Document, portfolio_data):
    """
    Inserts one or more Investment Holdings tables into the 'document'.

    :param document: The docx Document object where tables should be inserted.
    :param portfolio_data: Either a single JSON object (dict) or a list of JSON objects.
        Each JSON should look like:
        {
          "PortfolioTotal": <number>,
          "Holdings": [
             {"Fund": "<Fund Name>", "Value": <number>, "Percent": "<percentage string>"},
             ...
          ]
        }
    """

    # If multiple portfolio JSON objects exist, handle them in a loop
    if isinstance(portfolio_data, list):
        # Insert each portfolio's table
        for idx, single_portfolio in enumerate(portfolio_data, start=1):
            holdings = single_portfolio.get("Holdings", [])
            # Add a heading or some label to indicate which file this table belongs to
            document.add_heading(f"Investment Holdings (File {idx})", level=2)

            # Create the table in the same Document
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            # Fill header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Fund"
            hdr_cells[1].text = "Value"
            hdr_cells[2].text = "Percent"

            # Fill rows
            for h in holdings:
                row_cells = table.add_row().cells
                row_cells[0].text = str(h.get("Fund", ""))
                row_cells[1].text = str(h.get("Value", ""))
                row_cells[2].text = str(h.get("Percent", ""))

            document.add_paragraph("")  # extra spacing
    else:
        # We only have a single portfolio JSON
        single_portfolio = portfolio_data
        holdings = single_portfolio.get("Holdings", [])

        document.add_heading("Investment Holdings", level=2)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Fill header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Fund"
        hdr_cells[1].text = "Value"
        hdr_cells[2].text = "Percent"

        # Fill rows
        for h in holdings:
            row_cells = table.add_row().cells
            row_cells[0].text = str(h.get("Fund", ""))
            row_cells[1].text = str(h.get("Value", ""))
            row_cells[2].text = str(h.get("Percent", ""))

        document.add_paragraph("")



"""
def create_annuity_quotes_table(document, annuity_quotes):
  
    if not annuity_quotes or "Quotes" not in annuity_quotes:
        raise ValueError("Invalid annuity quotes data.")

    quotes = annuity_quotes["Quotes"]

    # Add a heading for the table
    document.add_heading("Annuity Quotes", level=2)

    # Create the table with rows for attributes and a header for quotes
    table = document.add_table(rows=5, cols=len(quotes) + 1)
    table.style = 'Table Grid'

    # Fill the first cell in the header with an empty label
    table.cell(0, 0).text = ""

    # Add headers for Quotes (Quote 1, Quote 2, etc.)
    for idx in range(len(quotes)):
        table.cell(0, idx + 1).text = f"Quote {idx + 1}"

    # Define the rows for attributes
    attributes = ["Purchase Amount", "Monthly Amount", "Yearly Amount", "Yearly Increase"]

    for row_idx, attribute in enumerate(attributes, start=1):
        # Add the attribute name to the first column
        table.cell(row_idx, 0).text = attribute
        # Fill in the data for each quote
        for col_idx, quote in enumerate(quotes):
            table.cell(row_idx, col_idx + 1).text = quote.get(attribute, "")

    return table
    
    """
 

def create_new_document(template_path, factfinding_text, plan_review_paragraphs, portfolio_json, attitude_to_risk,
                        table_data, product_report_text, plan_report_text, last_year_performance_text,
                        fund_performance_data, dark_star_performance_data, sap_comparison_tables,
                        annuity_quotes_text, fund_comparison_text, plan_review_texts,
                        safe_withdrawal_text,iht_text, output_path):
    """
    Create a well-formatted document by replacing placeholders, appending tables,
    and inserting dynamically generated sections while preserving static text.
    """

 # Parse client details from GPT response
    client_details_json_str = extract_client_details_with_gpt(factfinding_text)
    try:
        client_details = extract_client_details_with_gpt(factfinding_text)
    except json.JSONDecodeError as e:
        raise ValueError(
            "Error parsing client details JSON: " + repr(e) +
            "\nRaw GPT Response:\n" + client_details_json_str
        )
    """# Validate risk details
    if not isinstance(risk_details, dict):
        raise ValueError(f"Expected risk_details to be a dictionary, but got: {type(risk_details)}")"""

    # Generate dynamic sections
    current_situation = generate_current_situation(factfinding_text)
    priorities_and_objectives = generate_priorities_and_objectives(factfinding_text)

    pension_review_section = generate_pension_review_section(product_report_text)
    swr_section = safe_withdrawal_text
    

    # Load the template
    original_doc = Document(template_path)
    new_doc = Document()

  
    # Replace placeholders while preserving static text
    for paragraph in original_doc.paragraphs:
        text = paragraph.text

        # Replace placeholders dynamically
        if "{Full name}" in text:
            text = text.replace("{Full name}", client_details.get("Full name", ""))
        if "{Address}" in text:
            text = text.replace("{Address}", client_details.get("Address", ""))
        if "{Today’s date}" in text:
            text = text.replace("{Today’s date}", client_details.get("Today’s date", ""))
        if "{salutation}" in text:
            text = text.replace("{salutation}", client_details.get("salutation", ""))
        if "{Current_Situation}" in text:
            text = text.replace("{Current_Situation}", current_situation)
        if "{Priorities_and_Objectives}" in text:
            text = text.replace("{Priorities_and_Objectives}", priorities_and_objectives)
        if "{Attitude_to_Risk}" in text:
            text = text.replace("{Attitude_to_Risk}", attitude_to_risk or "")
        if "{Review of Existing Royal London Personal Pension}" in text:
            text = text.replace("{Review of Existing Royal London Personal Pension}", "")  # Remove placeholder
            new_doc.add_heading("Plan Reviews", level=2)
            for review in plan_review_paragraphs:
                new_doc.add_paragraph(review)


        if "{Safe Withdrawal Rate (SWR)}" in text:
            text = text.replace("{Safe Withdrawal Rate (SWR)}", swr_section)

        # Handle the {table} placeholder
        if "{table1}" in text:
            paragraph.text = "Overview of Current Investments\nHere is a breakdown of your current investments."
            create_plan_report_table(new_doc, table_data)  # Insert table right after the paragraph
            continue
        # Handle the {table2-1} placeholder for fund performance

        # Example placeholder for last-year performance
        if "{Last_Year_Performance}" in text:
            # If last_year_performance_text is provided, use it
            if last_year_performance_text:
                text = text.replace("{Last_Year_Performance}", last_year_performance_text)
            else:
                text = text.replace("{Last_Year_Performance}", "No single-year performance data found.")
        # Check for the new Investment_holdings placeholder

        if "{Investment_holdings}" in text:
            # Remove the placeholder from the paragraph text
            text = text.replace("{Investment_holdings}", "")
            # Insert the tables for one or more portfolios
            if portfolio_json:
                add_investment_holdings_tables(new_doc, portfolio_json)
            else:
                new_doc.add_paragraph("No portfolio data found for Investment Holdings.")


        if "{table2-1}" in text:
            bullet_points = "Extracted Fund Performance\n\n"
            # Ensure fund_performance_data is a list so we can iterate over it
            if not isinstance(fund_performance_data, list):
                fund_performance_data = [fund_performance_data]
            for fund in fund_performance_data:
                bullet_points += f"**{fund['Fund']}**\n"
                for year in range(1, 6):
                    year_key = f"Year {year}"
                    benchmark_key = fund.get("Benchmark", {}).get(year_key, "N/A")
                    year_value = fund.get(year_key, "N/A")
                    bullet_points += f"- {year_key}: {year_value} (Benchmark: {benchmark_key})\n"
                cumulative_performance = fund.get("Cumulative (5 YR)", "N/A")
                cumulative_benchmark = fund.get("Benchmark", {}).get("Cumulative (5 YR)", "N/A")
                bullet_points += f"- Cumulative 5-Year Performance: {cumulative_performance} (Benchmark: {cumulative_benchmark})\n\n"
            text = text.replace("{table2-1}", bullet_points.strip())


                # Handle the {table2-2} placeholder for Dark Star performance
        if "{table2-2}" in text:
            bullet_points = "Extracted Dark Star Performance\n\n"
            # Ensure dark_star_performance_data is iterable (a list)
            if not isinstance(dark_star_performance_data, list):
                dark_star_performance_data = [dark_star_performance_data]
            for fund in dark_star_performance_data:
                bullet_points += f"**{fund.get('Fund', 'Unknown Fund')}**\n"
                for year in range(1, 6):
                    year_key = f"Year {year}"
                    year_value = fund.get(year_key, "N/A")
                    bullet_points += f"- {year_key}: {year_value}\n"
                cumulative_performance = fund.get("Cumulative (5 YR)", "N/A")
                bullet_points += f"- Cumulative 5-Year Performance: {cumulative_performance}\n\n"
            text = text.replace("{table2-2}", bullet_points.strip())


                # Handle placeholders
        # Handle the {table3-1} placeholder for SAP comparison tables (multiple files)
        # Handle the {table3-1} placeholder for SAP comparison tables (multiple files)
        if "{table3-1}" in text:
            text = text.replace("{table3-1}", "")
            if sap_comparison_tables:
                for sc_table in sap_comparison_tables:
                    age = sc_table.get("Age", "N/A")
                    # Search for a row key containing "Effect on Fund if Moved"
                    effect_key = None
                    for key in sc_table["Table"].keys():
                        if "Effect on Fund if Moved" in key:
                            effect_key = key
                            break
                    if effect_key:
                        effect_values = sc_table["Table"].get(effect_key, ["N/A", "N/A", "N/A"])
                        if len(effect_values) != 3:
                            effect_values = ["N/A", "N/A", "N/A"]
                        try:
                            middle_value_str = effect_values[1]
                            middle_value = float(middle_value_str.strip('%').strip()) if middle_value_str != "N/A" else 0.0
                        except ValueError:
                            middle_value = 0.0
                    else:
                        middle_value_str = "N/A"
                        middle_value = 0.0

                    # Build explanatory text based on the middle value
                    if middle_value < 0.0:
                        below_table_text = (
                            f"The critical yield required to match the benefits of your current scheme at age {age} "
                            f"is {middle_value_str}, indicating that the proposed arrangement would need less performance "
                            f"per annum to make up the costs of transferring. This is because the proposed arrangement "
                            "is more cost-effective than your current arrangement."
                        )
                    elif 0.0 <= middle_value < 3.0:
                        below_table_text = (
                            f"The critical yield required to match the benefits of your current scheme at age {age} "
                            f"is {middle_value_str}, indicating that the proposed arrangement would need an additional "
                            f"fund performance per annum to make up the costs of transferring.\n\n"
                            "I believe the chosen fund will be able to achieve this over the long term, although this is not guaranteed."
                        )
                    else:
                        below_table_text = (
                            f"The critical yield required to match the benefits of your current scheme at age {age} "
                            f"is {middle_value_str}, indicating that the proposed arrangement would need an additional "
                            f"performance per annum to make up the costs of transferring.\n\n"
                            "I cannot guarantee that the recommended fund can match the additional performance required to make up the costs of transferring, "
                            "but I still believe that transferring out is in your best interests. Performance is only one consideration to make when transferring out."
                        )

                    heading = f"Comparison at Age {age}"
                    paragraph_before_table = (
                        f"The table below shows the projected value of your pensions at the age of {age}, firstly if it were to remain in your current arrangement and secondly were it to be transferred."
                    )
                    new_doc.add_heading(heading, level=2)
                    new_doc.add_paragraph(paragraph_before_table)
                    create_comparison_table(new_doc, sc_table)
                    new_doc.add_paragraph("")  # Blank line after the table
                    new_doc.add_paragraph(below_table_text)
            

        # Insert Annuity Quotes text (if any)
        if "{Annuity_Quotes}" in text:
            if annuity_quotes_text:
                text = text.replace("{Annuity_Quotes}", annuity_quotes_text)
            else:
                text = text.replace("{Annuity_Quotes}", "No annuity quotes available.")

                # Replace placeholder for Fund Comparison
        if "{Fund_Comparison}" in text:
            text = text.replace("{Fund_Comparison}", fund_comparison_text)  

# Handle the {IHT_Table} placeholder with bullet points
        if "{IHT_Text}" in text:
            text = text.replace("{IHT_Text}", iht_text)
        # Add formatted paragraph to the new document
        new_paragraph = new_doc.add_paragraph(text)
        new_paragraph.style = new_doc.styles['Normal']
        new_paragraph.paragraph_format.space_before = Pt(6)
        new_paragraph.paragraph_format.space_after = Pt(6)
        new_paragraph.paragraph_format.line_spacing = Pt(12)  # Set consistent line spacing
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Adjust document margins
    section = new_doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Save the combined document
    new_doc.save(output_path)
    print(f"Debug: Document saved successfully at {output_path}")


    return swr_section


def save_uploaded_file(uploaded_file, folder):
    import os
    os.makedirs(folder, exist_ok=True)
    file_path = os.path.join(folder, uploaded_file.name)
    
    # Make absolutely sure we write the entire buffer
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())  # <-- getbuffer instead of read()

    return file_path

